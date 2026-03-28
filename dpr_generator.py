"""
dpr_generator.py  –  Generates full DPR .docx matching the approved template.
Accepts a params dict, returns bytes.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
import io, os, copy
from financial_model import run_model

# ── Ownership string helper ────────────────────────────────────────────────────
def ownership_str(p):
    """Return 'Parent1 (X%) & Parent2 (Y%)' or just 'Parent1 (100%)' for sole owners."""
    if p.get("parent2_name") and p.get("parent2_pct", 0) > 0:
        return f"{p['parent1_name']} ({p['parent1_pct']:.0f}%) & {p['parent2_name']} ({p['parent2_pct']:.0f}%)"
    return f"{p['parent1_name']} ({p['parent1_pct']:.0f}%)"

def has_parent2(p):
    return bool(p.get("parent2_name")) and p.get("parent2_pct", 0) > 0

# ── Image paths ──────────────────────────────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))

# Bundled default images (fallback when no custom image is uploaded)
_DEFAULT_IMAGES = {
    "img_cover":          os.path.join(_HERE, "static", "images", "img_cover.jpeg"),
    "img_basic_concept":  os.path.join(_HERE, "static", "images", "img_basic_concept.jpeg"),
    "img_solar_radiation":os.path.join(_HERE, "static", "images", "img_solar_radiation.jpeg"),
    "img_location":       os.path.join(_HERE, "static", "images", "img_location_map.png"),
    "img_vicinity":       os.path.join(_HERE, "static", "images", "img_vicinity_map.png"),
    "img_district":       os.path.join(_HERE, "static", "images", "img_district_map.jpeg"),
}

# Legacy alias so existing call sites still work
IMAGES = {
    "cover":           _DEFAULT_IMAGES["img_cover"],
    "basic_concept":   _DEFAULT_IMAGES["img_basic_concept"],
    "solar_radiation": _DEFAULT_IMAGES["img_solar_radiation"],
    "location_map":    _DEFAULT_IMAGES["img_location"],
    "vicinity_map":    _DEFAULT_IMAGES["img_vicinity"],
    "district_map":    _DEFAULT_IMAGES["img_district"],
}

# ── Colours (matching template) ───────────────────────────────────────────────
C_NAVY   = RGBColor(0x1F, 0x38, 0x64)
C_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
C_ORANGE = RGBColor(0xC5, 0x5A, 0x11)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN  = RGBColor(0x1E, 0x7B, 0x34)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_GOLD   = RGBColor(0xD4, 0xAF, 0x37)

# ── Number helpers ────────────────────────────────────────────────────────────
def f2(n):  return f"{n:,.2f}"
def f0(n):  return f"{n:,.0f}"

# ── XML / cell helpers ────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def set_cell_border(cell, color="AAAAAA"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0");    b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)

def set_col_width(cell, width_cm):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

# ── Paragraph helpers ─────────────────────────────────────────────────────────
def normal_para(doc, text, size_pt=None, bold=False, italic=False, color=None,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after_pt=6, space_before_pt=3, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.underline = underline
    r.font.name = "Arial"
    if size_pt:  r.font.size = Pt(size_pt)
    if color:    r.font.color.rgb = color
    return p

def caption(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    return p

def bullet(doc, text, size_pt=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = "Arial"
    if size_pt: r.font.size = Pt(size_pt)
    return p

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(9)
    for r in p.runs:
        r.font.name = "Arial"; r.font.size = Pt(15)
        r.font.color.rgb = C_NAVY; r.bold = True
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after  = Pt(6)
    for r in p.runs:
        r.font.name = "Arial"; r.font.size = Pt(13)
        r.font.color.rgb = C_BLUE; r.bold = True
    return p

def h3(doc, text, size_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size_pt)
    r.font.name = "Arial"; r.font.color.rgb = C_NAVY
    return p

def page_break(doc):
    doc.add_page_break()

# Mapping: legacy IMAGES key  ->  params key  ->  default key
_IMG_PARAMS_KEY = {
    "cover":           "img_cover",
    "basic_concept":   "img_basic_concept",
    "solar_radiation": "img_solar_radiation",
    "location_map":    "img_location",
    "vicinity_map":    "img_vicinity",
    "district_map":    "img_district",
}

def add_image(doc, key, width_cm, height_cm=None, cap_text=None, params=None):
    """Insert an image. Prefers custom upload from params, falls back to bundled default."""
    path = None

    # 1. Check if a custom upload was provided via params
    if params:
        param_key = _IMG_PARAMS_KEY.get(key)
        if param_key:
            uploaded = params.get(param_key)
            if uploaded and os.path.exists(str(uploaded)):
                path = str(uploaded)

    # 2. Fall back to bundled default
    if not path:
        default_key = _IMG_PARAMS_KEY.get(key)
        path = _DEFAULT_IMAGES.get(default_key, IMAGES.get(key, ""))

    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run()
        if height_cm:
            r.add_picture(path, width=Cm(width_cm), height=Cm(height_cm))
        else:
            r.add_picture(path, width=Cm(width_cm))
    else:
        p = normal_para(doc, f"[Image: {key} - not found]",
                        italic=True, color=RGBColor(0x99,0x99,0x99), size_pt=9)
    if cap_text:
        caption(doc, cap_text)
    return p

# ── Table builder ─────────────────────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths_cm,
               header_bg="1F3864", alt_bg="EBF3FB",
               header_size=9, body_size=9,
               center_cols=None):
    """
    center_cols: list of col indices to centre (default: all)
    """
    ncols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=ncols)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "888888")
        set_col_width(cell, col_widths_cm[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h); r.bold = True
        r.font.size = Pt(header_size)
        r.font.name = "Arial"; r.font.color.rgb = C_WHITE

    # Data rows
    for ri, row in enumerate(rows):
        bg = alt_bg if ri % 2 == 1 else "FFFFFF"
        # Check for section header rows (only 1 meaningful cell)
        is_section = (len([v for v in row if str(v).strip()]) <= 2
                      and ri > 0
                      and str(row[0]).strip() == ""
                      and str(row[1] if len(row)>1 else "").isupper())
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            use_bg = "D6E4F0" if is_section else bg
            set_cell_bg(cell, use_bg)
            set_cell_border(cell, "AAAAAA")
            set_col_width(cell, col_widths_cm[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            # Alignment
            if center_cols is None or ci in (center_cols or []):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(body_size)
            r.font.name = "Arial"
            r.bold = is_section

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ════════════════════════════════════════════════════════════════════════════
# MAIN GENERATOR
# ════════════════════════════════════════════════════════════════════════════
def generate_dpr(params: dict) -> bytes:
    m        = run_model(params)
    cfs      = m["cashflows"]
    proj_irr = m["project_irr"]
    eq_irr   = m["equity_irr"]
    min_dscr = m["min_dscr"]
    avg_dscr = m["avg_dscr"]
    ann_prin = m["ann_principal"]
    dep_slm  = m["dep_slm"]
    p        = params

    # ── Derived shorthand ────────────────────────────────────────────────────
    cap_ac  = p["project_capacity_ac"]
    cap_dc  = p["project_capacity_dc"]
    gen_yr1 = p["gen_yr1_lac_per_mwac"] * cap_ac
    loc_str = f"{p['location_village']}, {p['location_taluk']}, {p['location_district']}, {p['location_state']}"

    doc = Document()

    # ── Default Normal style ──────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(6)
    doc.styles["Normal"].paragraph_format.space_before = Pt(3)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 0  –  Portrait (cover → Section 12)
    # ════════════════════════════════════════════════════════════════════════
    sec0 = doc.sections[0]
    sec0.page_width    = Cm(21.59); sec0.page_height   = Cm(27.94)
    sec0.left_margin   = Cm(1.91);  sec0.right_margin  = Cm(1.91)
    sec0.top_margin    = Cm(1.91);  sec0.bottom_margin = Cm(1.91)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    def cline(text, sz, bold=False, col=C_NAVY, sa=4, sb=0, align=WD_ALIGN_PARAGRAPH.CENTER):
        pp = doc.add_paragraph()
        pp.alignment = align
        pp.paragraph_format.space_after  = Pt(sa)
        pp.paragraph_format.space_before = Pt(sb)
        r = pp.add_run(text)
        r.bold = bold; r.font.size = Pt(sz)
        r.font.name = "Arial"; r.font.color.rgb = col
        return pp

    cline("DETAILED PROJECT REPORT (DPR)", 20, bold=True, col=C_NAVY, sa=12, sb=24)
    add_image(doc, "cover", width_cm=15.9, height_cm=10.1, params=p)
    cline(f"DEVELOPMENT OF {f2(cap_dc)} MWp DC / {f2(cap_ac)} MW AC", 17, bold=True, sa=6, sb=6)
    cline("GROUND MOUNTED SOLAR PV POWER PROJECT", 14, bold=True, sa=4)
    cline(f"AT {p['location_village'].upper()} VILLAGE", 13, bold=True, col=C_ORANGE, sa=4)
    cline(f"{p['location_district'].upper()} DISTRICT, {p['location_state'].upper()}", 13, bold=True, col=C_ORANGE, sa=4)
    cline("Prepared for:", 10, col=C_BLUE, sa=5, sb=12)
    cline(f"M/S. {p['company_name'].upper()} ({p['company_short'].upper()})", 14, bold=True, sa=4)
    cline(p["company_address"], 10, col=RGBColor(0x44,0x44,0x44), sa=3)
    cline(f"A Subsidiary of {ownership_str(p)}", 10, col=RGBColor(0x55,0x55,0x55), sa=3)
    cline(f"COD: {p['cod_month']} {p['cod_year']}", 12, bold=True, col=C_ORANGE, sa=3, sb=6)
    page_break(doc)

    # ── INDEX ────────────────────────────────────────────────────────────────
    normal_para(doc, "INDEX", size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_after_pt=10, space_before_pt=12)
    for item in [
        "1.  Summary of the Company & Group Profile",
        "2.  Proposal for Solar Power Project",
        "3.  Concept of Captive Power Sale",
        "4.  Solar Resource Potential & Radiation Terminology",
        "5.  Project Site Solar Resource Assessment",
        "6.  Project at a Glance",
        "7.  Demand Analysis and Justification",
        "8.  Benefits of Grid Connected Solar PV Power Plant",
        "9.  Basic System Description",
        "10.  Bill of Quantity",
        "11.  Planned Project Schedule",
        "12.  Project Costing",
        "13.  Project Financials – 25 Year Cash Flow",
        "14.  Debt Repayment Schedule & DSCR",
        "15.  Summary of Results",
        "16.  Conclusion",
    ]:
        normal_para(doc, item, space_after_pt=4, space_before_pt=0,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
    page_break(doc)

    # ── SECTION 1 – COMPANY ──────────────────────────────────────────────────
    h1(doc, "1. SUMMARY OF THE COMPANY")
    if has_parent2(p):
        normal_para(doc, f"{p['company_name']} ({p['company_short']}) is a renewable energy focused Independent Power "
                    f"Producer (IPP), incorporated as a joint venture subsidiary with equal participation from:",
                    space_after_pt=6, space_before_pt=3)
        bullet(doc, f"{p['parent1_name']} – {p['parent1_pct']:.0f}%")
        bullet(doc, f"{p['parent2_name']} – {p['parent2_pct']:.0f}%")
    else:
        normal_para(doc, f"{p['company_name']} ({p['company_short']}) is a renewable energy focused Independent Power "
                    f"Producer (IPP), incorporated as a wholly owned subsidiary of {p['parent1_name']}:",
                    space_after_pt=6, space_before_pt=3)
        bullet(doc, f"{p['parent1_name']} – 100%")
    normal_para(doc, f"{p['company_short']} is registered at {p['company_address']}. The company is established with a "
                "clear vision to harness renewable energy and cater to the growing energy needs of industrial and "
                f"commercial consumers in {p['location_state']}.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, f"{p['company_short']} intends to invest in a {f2(cap_ac)} MW AC ({f2(cap_dc)} MWp DC) "
                f"ground-mounted Solar Power Project at {p['location_village']} Village, {p['location_district']} "
                f"District, {p['location_state']}. The power generated from the solar plant will be sold under a "
                f"{p['ppa_term']}-year PPA at a flat tariff of ₹ {p['ppa_tariff']:.2f} per unit. "
                f"COD is planned for {p['cod_month']} {p['cod_year']}.", space_after_pt=6, space_before_pt=3)
    if has_parent2(p):
        normal_para(doc, f"{p['company_short']} combines the strengths of its two parent organisations – "
                    f"{p['parent1_name']}'s deep engineering expertise, project execution capabilities, and financial "
                    f"strength, with {p['parent2_name']}'s specialised renewable energy project development, EPC "
                    "capabilities, and operational experience in solar and wind power assets.",
                    space_after_pt=6, space_before_pt=3)
    else:
        normal_para(doc, f"{p['company_short']} leverages {p['parent1_name']}'s deep engineering expertise, "
                    "project execution capabilities, financial strength, and specialised renewable energy project "
                    "development experience to deliver this project.",
                    space_after_pt=6, space_before_pt=3)

    h2(doc, "Group Profile")
    if has_parent2(p):
        normal_para(doc, "The Group operates through two synergistic parent entities that collectively bring more than "
                    "three decades of combined operating experience in engineering, automation, and renewable energy sectors:",
                    space_after_pt=4, space_before_pt=3)
    else:
        normal_para(doc, f"The project is promoted by {p['parent1_name']}, which brings significant engineering, "
                    "automation, and project execution experience across multiple sectors:",
                    space_after_pt=4, space_before_pt=3)

    h3(doc, f"I. {p['parent1_name']}")
    normal_para(doc, f"{p['parent1_name']} is one of the promoter entities holding {p['parent1_pct']:.0f}% equity in "
                f"{p['company_short']}. The company brings significant financial strength, industry relationships, "
                "and execution capabilities. It has a proven track record in its domain with experienced management "
                "and strong corporate governance frameworks.", space_after_pt=5, space_before_pt=3)

    if has_parent2(p):
        h3(doc, f"II. {p['parent2_name']}")
        normal_para(doc, f"{p['parent2_name']} holds {p['parent2_pct']:.0f}% equity in {p['company_short']}. "
                    "The company contributes renewable energy project development expertise, EPC execution capability, "
                    "and operational experience in solar and wind power assets across India. Its integrated model covering "
                    "project development, EPC execution, asset ownership, and O&M ensures strong control over cost, "
                    "quality, and long-term performance.", space_after_pt=5, space_before_pt=3)

    h2(doc, "Group Vision & Mission")
    normal_para(doc, "Group Vision: To build a globally respected engineering and clean energy group delivering "
                "reliable, sustainable, and technology-driven solutions that support the world's transition to a "
                "low-carbon future.", space_after_pt=4, space_before_pt=3)
    normal_para(doc, "Group Mission: To create long-term value by combining engineering excellence, innovation, and "
                "disciplined execution across automation and renewable energy businesses, delivering efficient, "
                "reliable, and sustainable solutions for customers worldwide.", space_after_pt=6, space_before_pt=3)

    h2(doc, f"About the Project Company – {p['company_short']}")
    normal_para(doc, f"{p['company_name']} ({p['company_short']}) has been incorporated specifically to develop, "
                "construct, own, and operate the proposed solar power project. As a special purpose vehicle (SPV), "
                f"{p['company_short']} will have dedicated focus on this project, with ring-fenced assets and "
                f"liabilities. The company is owned by {ownership_str(p)}.", space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"{p['company_short']} will be responsible for: obtaining all necessary approvals, permits, "
                "and consents; entering into the Power Purchase Agreement with the identified consumer; procuring "
                "financing from lenders; appointing and supervising the EPC contractor; and managing the plant "
                "during the operational phase either directly or through an appointed O&M contractor.",
                space_after_pt=6, space_before_pt=3)
    page_break(doc)

    # ── SECTION 2 – PROPOSAL ─────────────────────────────────────────────────
    h1(doc, "2. PROPOSAL FOR SOLAR POWER PROJECT")
    normal_para(doc, f"The company is setting up this project considering the attractive opportunities in the "
                f"non-conventional energy sector aided by the Solar Energy Policy of the Government of "
                f"{p['location_state']}. M/s. {p['company_name'].upper()} ({p['company_short']}) is proposing to "
                f"set up a {f2(cap_ac)} MW AC ({f2(cap_dc)} MWp DC) Ground Mounted Solar PV Power Project at "
                f"{p['location_village'].upper()} Village, {p['location_district'].upper()} District, "
                f"{p['location_state']}.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, f"The annual power generation from the {f2(cap_ac)} MW AC, {f2(cap_dc)} MWp DC Solar Plant "
                f"will be approximately {f2(gen_yr1)} Lac Units ({gen_yr1/10:.1f} Million kWh) per annum in the "
                f"first year. The power generated from the solar plant will be sold under a {p['ppa_term']}-year "
                f"Power Purchase Agreement (PPA) at a flat tariff of ₹ {p['ppa_tariff']:.2f} per kWh for "
                f"{p['ppa_term']} years.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, f"India's renewable energy sector has witnessed unprecedented growth in the past decade, "
                "driven by Government policies, falling equipment costs, and increasing awareness about "
                "sustainability. The Government of India has set an ambitious renewable energy target of 500 GW "
                "by 2030, out of which solar energy alone is expected to contribute 300 GW.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, f"The proposed project by {p['company_short']} is strategically located in "
                f"{p['location_district']} District, which has one of the highest solar irradiation levels in "
                f"{p['location_state']}. The region receives an average of 5.5 to 5.8 kWh per square metre per "
                "day of Global Horizontal Irradiance (GHI), making it highly suitable for solar power generation.",
                space_after_pt=6, space_before_pt=3)
    normal_para(doc, "The project is expected to generate significant employment during construction and operation "
                "phases, contribute to reduction in carbon emissions, and supply clean, affordable electricity.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "Project Rationale")
    normal_para(doc, "The rationale for setting up this solar power project is based on the following key drivers:",
                space_after_pt=4, space_before_pt=3)
    for item in [
        f"Rising electricity demand in {p['location_state']}'s industrial and commercial sectors, with an average annual demand growth of 6–8%",
        "Increasing cost of conventional grid power, making long-term fixed-tariff renewable energy a cost-effective alternative",
        f"Availability of favourable solar resources at the proposed project site in {p['location_village']}",
        f"Strong policy support from the Government of {p['location_state']} and MNRE, including open access, net metering, and wheeling provisions",
        "Declining capital costs of solar PV equipment, improving project economics",
        f"Long-term PPA at ₹ {p['ppa_tariff']:.2f} per unit providing revenue certainty and bankability",
        f"Strong promoter group credentials with {ownership_str(p)} as parent {'companies' if has_parent2(p) else 'company'}",
    ]:
        bullet(doc, item)

    h2(doc, "Expected Impact of the Project")
    normal_para(doc, f"The commissioning of the {f2(cap_ac)} MW solar power project by {p['company_short']} at "
                f"{p['location_village']} will have significant positive impacts across multiple dimensions:",
                space_after_pt=4, space_before_pt=3)
    for item in [
        f"Generation: ~{f2(gen_yr1)} Lac Units/year in Year 1; ~{gen_yr1*23:.0f} Lac Units over 25 years",
        "Employment: ~200 jobs during construction; ~15 permanent O&M jobs",
        f"CO₂ Avoidance: ~{gen_yr1*0.82*100:,.0f} tonnes/year (~{gen_yr1*0.82*100*25/100000:.1f} lakh tonnes over 25 years)",
        f"Revenue for Promoters: Total projected revenue over 25 years at ₹ {p['ppa_tariff']:.2f} per unit is approximately ₹ {gen_yr1*p['ppa_tariff']*22/100:,.0f} Crores",
        f"Local Development: Land lease payments and local procurement will benefit the {p['location_district']} District economy",
    ]:
        bullet(doc, item)
    page_break(doc)

    # ── SECTION 3 – CONCEPT OF CAPTIVE POWER SALE ────────────────────────────
    h1(doc, "3. CONCEPT OF CAPTIVE POWER SALE")
    normal_para(doc, f"The energy generated from the project shall be fully sold to the procurer under a "
                f"{p['ppa_term']}-year flat-rate Power Purchase Agreement (PPA) at ₹ {p['ppa_tariff']:.2f} per unit. "
                "The concept of power sale through PPA is a well-established mechanism in the Indian renewable "
                "energy sector, enabling generators and consumers to enter into long-term, predictable energy "
                "supply arrangements.", space_after_pt=6, space_before_pt=3)

    h2(doc, "Basic Concept of Power Sale")
    normal_para(doc, "Under the PPA model, the solar generation plant supplies power into the nearest power grid, "
                "which is then credited against the energy consumption of the identified consumer locations. "
                "The basic concept is illustrated in the diagram below:", space_after_pt=5, space_before_pt=3)
    add_image(doc, "basic_concept", width_cm=14.8, height_cm=13.0, params=p)
    caption(doc, "Figure: Basic Concept of Solar Power Sale / Captive Wheeling")
    normal_para(doc, "As illustrated in the diagram above, the Existing Power Supply Network (ESCOM) currently "
                "supplies power to multiple consumer locations (Consumer Location 1, 2, 3, etc.), each metered "
                "separately. The total consumption is the aggregate of all consumer meter readings.",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"Under the proposed arrangement, the Solar Generation Plant supplies energy into the Nearest "
                "Power Grid, which is metered (Meter Reading 4). This energy is then deducted from the EB bills of "
                "the identified consumers, effectively adjusting the energy cost against the solar generator. "
                f"The net result is that the consumers receive solar-generated power at ₹ {p['ppa_tariff']:.2f} "
                "per unit – significantly lower than the prevailing HT industrial tariff.",
                space_after_pt=6, space_before_pt=3)

    h3(doc, "Advantages of Solar Power PPA", size_pt=12)
    for item in [
        "A viable alternative solution to affordable power to the consumers on a long-term basis at competitive rates",
        "Increases the competitiveness of the consumer company by reducing energy cost burden",
        f"Huge savings in energy costs on long-term basis – estimated savings of ₹ 1.50 to ₹ 2.50 per unit against prevailing grid tariff",
        "Self-sufficient and Green Energy for the consumer through solar power production",
        f"Hedge against rising grid electricity prices – flat PPA tariff for {p['ppa_term']} years insulates against tariff escalation",
        "Reduced carbon footprint and ESG (Environmental, Social, Governance) compliance",
        "Supports the consumer's sustainability reporting, green certification, and corporate responsibility initiatives",
        "No capital investment required from the consumer – the generator bears all project costs",
    ]:
        bullet(doc, item)

    h2(doc, "Policy Framework")
    normal_para(doc, f"The project is being developed under the enabling policy framework of the Government of "
                f"{p['location_state']} and the Government of India. Key policy initiatives that support this "
                "project include:", space_after_pt=4, space_before_pt=3)
    for item in [
        f"{p['location_state']} Solar Energy Policy 2023 – Targets 9,000 MW of solar capacity by 2025",
        "National Solar Mission under the National Action Plan on Climate Change (NAPCC)",
        "Open Access provisions under the Electricity Act 2003, enabling wheeling of power through the grid",
        "Exemption from payment of Electricity Tax up to 100% for 5 years for eligible solar projects",
        "Cross-subsidy surcharge exemptions for captive/group captive solar projects",
        "Provision for banking of energy units generated but not immediately consumed",
        f"Wheeling charges as applicable under {p['location_state']} Electricity Regulatory Commission regulations",
    ]:
        bullet(doc, item)

    normal_para(doc, f"The Government of {p['location_state']} is committed to mitigating climate change effects "
                "by bringing out policies conducive to promoting renewable energy generation in the state. "
                f"{p['location_state']} Solar Energy Policy 2023 aims at generating 9,000 MW of solar power by "
                "2025. India expanded its solar generation capacity many times from 2,650 MW on 26th May 2014 "
                "to over 70 GW as on 31st March 2023.", space_after_pt=6, space_before_pt=6)

    h2(doc, "Background of the Project")
    normal_para(doc, "Large multi-megawatt PV plants, of more than 500 GW in aggregate, are now in operation "
                "worldwide. Indian solar power installations as on 31st March 2023 touched 70 GW. Solar "
                "Photovoltaic (PV) is known to be an important energy source for developing countries like India. "
                "Its importance is now being reaffirmed even by developed countries in view of its renewable and "
                "environment-friendly character.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, "In India, solar PV is now commercially operated by independent power producers. The Ministry "
                "of New & Renewable Energy Sources has been promoting electricity generation from Solar PV at the "
                "Mega-Watt level under the Grid Interactive Solar PV Power Generation Projects scheme. Pursuant "
                f"to this liberalisation policy, M/s. {p['company_name']} has proposed to set up a Solar Power "
                f"Project of {f2(cap_ac)} MW AC capacity in {p['location_village']} Village in the state of "
                f"{p['location_state']}.", space_after_pt=6, space_before_pt=3)

    h2(doc, "Selection of Technology")
    normal_para(doc, "The generation of energy from mega projects is commercially viable with conventional "
                "technologies such as Nuclear and Thermal power. Non-conventional energy technologies for small "
                "and medium power plants have been proven over the past ten years for reliable operation in the "
                f"states of {p['location_state']}, Karnataka, Andhra Pradesh, and Rajasthan.",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, "Based on available satellite data and PVSyst simulations, the company has analysed the "
                "solar radiation and other parameters including solar horizon, temperature coefficients, dust "
                "soiling, and system losses at the proposed site to confirm technical feasibility. "
                "The use of monocrystalline silicon solar PV modules has been selected for the following reasons:",
                space_after_pt=4, space_before_pt=3)
    for item in [
        "Higher efficiency (21–23%) compared to polycrystalline modules, resulting in lower area requirement",
        f"Better performance at high temperatures, critical for {p['location_state']}'s climate",
        "Proven long-term degradation rates of less than 0.5% per year after the first year",
        "Availability of reliable suppliers with Tier-1 ratings from BNEF (Bloomberg New Energy Finance)",
        "Strong performance warranty (10 years product warranty; 25-year linear performance warranty)",
    ]:
        bullet(doc, item)
    page_break(doc)

    # ── SECTION 4 – SOLAR RESOURCE ───────────────────────────────────────────
    h1(doc, "4. SOLAR RESOURCE POTENTIAL")
    h2(doc, "India's Solar Energy Resource")
    normal_para(doc, "India is located in the sunny belt of the earth, receiving abundant radiant energy from "
                "the sun. Its equivalent energy potential is about 6,000 million GWh of energy per year. India "
                "being a tropical country is blessed with good sunshine over most parts with 250–300 clear sunny "
                "days per year. The annual global radiation varies from 1,600 to 2,200 kWh/m².",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, "According to the National Renewable Energy Laboratory (NREL) and the National Institute "
                "of Wind and Solar Energy (NIWE), India has a theoretical solar energy potential of over 748 GW "
                "at current land utilisation norms. The practical potential is still well over 500 GW, "
                "significantly higher than India's current installed solar capacity of ~70 GW (March 2023).",
                space_after_pt=6, space_before_pt=3)

    h2(doc, f"{p['location_state']}'s Solar Resource")
    normal_para(doc, f"{p['location_state']} is one of the states with the highest solar irradiation in India. "
                "The state receives an annual average Global Horizontal Irradiance (GHI) of 5.0 to 6.2 kWh/m²/day "
                f"depending on location. The {p['location_district']} region, where the proposed project is "
                "located, falls in a high-irradiation zone with average GHI of approximately 5.5 to 5.8 "
                "kWh/m²/day.", space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"{p['location_state']} has set an ambitious solar energy target under the Solar Energy "
                "Policy 2023. The state has already commissioned over 6,000 MW of solar capacity and continues "
                "to attract large-scale investments from domestic and international developers. The state's "
                "proactive policy framework, including open access, net metering, group captive provisions, and "
                "dedicated solar parks, has made it one of the most attractive destinations for solar investment.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "Radiation Terminology")
    for term, defn in [
        ("Irradiance (W/m²)", "Amount of radiant energy incident on a surface per unit area per unit time. Measured using a pyranometer."),
        ("Direct Solar Irradiance (DNI)", "Solar irradiance on a surface held perpendicular to sun rays, excluding diffuse radiation. Measured by a pyrheliometer."),
        ("Global Horizontal Irradiance (GHI)", "Total solar irradiance on a horizontal surface due to both direct sun rays and diffuse sky radiation. The most common measurement used in solar resource assessment."),
        ("Diffuse Solar Irradiance (DHI)", "Solar irradiance on a horizontal surface due to sky radiation only (excluding direct sun rays). Important in overcast or partly cloudy conditions."),
        (f"Peak Sun Hours (PSH)", f"Equivalent number of hours per day when solar irradiance averages 1,000 W/m². The {p['location_district']} region averages 5.5–5.8 PSH."),
        ("Plane of Array Irradiance (POA)", "Total irradiance incident on the surface of the solar panel, accounting for tilt, orientation, and shading. Used in energy yield simulations."),
        ("Temperature Coefficient", "The rate at which solar module efficiency decreases with increasing temperature. For monocrystalline modules, this is typically -0.35% to -0.40% per °C above 25°C (STC)."),
        ("Capacity Factor (CUF)", f"The ratio of actual annual energy generation to the maximum possible generation if the plant operated at full rated capacity for the entire year. For this project, CUF is estimated at ~24.5%."),
    ]:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_after  = Pt(5)
        pp.paragraph_format.space_before = Pt(3)
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = pp.add_run(f"{term}: "); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = "Arial"; r1.font.color.rgb = C_NAVY
        r2 = pp.add_run(defn)
        r2.font.size = Pt(10); r2.font.name = "Arial"

    add_image(doc, "solar_radiation", width_cm=15.3, height_cm=6.9, params=p)
    caption(doc, "Figure: Solar Radiation Map – Tamil Nadu Region")

    h2(doc, "Solar Energy Generation – Key Metrics for the Project Site")
    make_table(doc,
        ["Solar Resource Parameter", f"Value at {p['location_village']} Site"],
        [
            ["Average GHI (Annual)", "~5.5 – 5.8 kWh/m²/day"],
            ["Average DNI (Annual)", "~5.0 – 5.3 kWh/m²/day"],
            ["Peak Sun Hours (PSH)", "~5.5 – 5.8 hours/day"],
            ["Annual Sunshine Hours", "~2,700 – 2,900 hours/year"],
            ["Average Ambient Temperature", "28°C – 32°C (Annual Average)"],
            [f"Specific Yield (Year 1)", f"~{p['gen_yr1_lac_per_mwac']*100/cap_dc:.0f} kWh/kWp/year"],
            ["Plant Capacity Utilisation Factor (CUF)", "~24.5% (Year 1)"],
            ["Performance Ratio (PR)", "~78% – 80%"],
        ],
        col_widths_cm=[9, 7], center_cols=[1])

    normal_para(doc, f"Note: Detailed site-specific solar resource assessment, PVSyst simulation reports, and "
                "long-term irradiance data are provided in Annexure – A at the end of this report.",
                italic=True, color=RGBColor(0x55,0x55,0x55), space_after_pt=10, space_before_pt=3)
    page_break(doc)

    # ── SECTION 5 – SITE DETAILS ─────────────────────────────────────────────
    h1(doc, "5. PROJECT SITE DETAILS & SOLAR RESOURCE ASSESSMENT")
    h2(doc, "Site Information")
    normal_para(doc, f"The proposed location of the solar power plant is in {p['location_taluk']} Tahsil "
                f"(Latitude: {p['latitude']}; Longitude: {p['longitude']}) of {p['location_district'].upper()} "
                f"District, {p['location_state']}. The project site is approximately {p['nearest_town_km']:.0f} km "
                f"from {p['nearest_town']} Town and is well connected by road.",
                space_after_pt=5, space_before_pt=3)

    add_image(doc, "location_map", width_cm=15.3, height_cm=13.8, params=p)
    caption(doc, f"Figure 1.1: Location Map – {p['location_village']}, {p['location_district']} District")
    add_image(doc, "vicinity_map", width_cm=15.3, height_cm=11.9, params=p)
    caption(doc, "Figure 1.2: Project Site Vicinity Map")
    add_image(doc, "district_map", width_cm=15.3, height_cm=13.2, params=p)
    caption(doc, f"Figure 1.3: District Map of {p['location_district']} – Proposed Project Site")

    h2(doc, "Site Introduction")
    normal_para(doc, f"The proposed location of the Solar Power Plant based on SPV technology is in "
                f"{p['location_taluk'].upper()} Tahsil of {p['location_district'].upper()} District of "
                f"{p['location_state']} state. The project is situated at {p['location_village'].upper()} Village, "
                f"approximately {p['nearest_town_km']:.0f} km from {p['nearest_town']} town.",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"The land where the project is proposed consists of private land admeasuring approximately "
                f"{p['land_acres']:.0f} acres. The land is a plain, shadow-free site, highly suitable for a "
                "utility-scale ground-mounted solar project with minimal civil work requirements.",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"The location is well connected with the state and national highway network. The proposed "
                "location has shadow-free area (almost flat terrain) and is located at very close proximity "
                f"(approximately {p['tneb_distance_km']:.1f} km) to the Evacuation Point at the grid substation, "
                "from where the power generated can be fed to the grid.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "Site Details")
    make_table(doc,
        ["Parameter", "Details"],
        [
            ["Area Required", f"Approx. {p['land_acres']:.0f} Acres (Private Land)"],
            ["Latitude", p["latitude"]],
            ["Longitude", p["longitude"]],
            ["Summer Temperature", "32°C to 40°C"],
            ["Winter Temperature", "17°C to 30°C"],
            ["Average Annual GHI", "~5.5 to 5.8 kWh/m²/day"],
            ["Annual Sunshine Hours", "~2,800 hours/year"],
            ["Terrain", "Plain land, shadow-free"],
            ["Distance to Grid SS", f"~{p['tneb_distance_km']:.1f} km from project site"],
            ["Land Type", "Private land – plain soil"],
            ["Nearest Town", f"{p['nearest_town']} (~{p['nearest_town_km']:.0f} km)"],
        ],
        col_widths_cm=[5.5, 10], center_cols=[])

    h2(doc, "Approach to Site")
    normal_para(doc, f"The project is located in {p['location_district'].upper()} District of "
                f"{p['location_state']} State, near {p['location_village'].upper()} Village. "
                f"The project area is accessible from {p['nearest_town']} town. "
                f"The distance from {p['location_village']} Village to {p['nearest_town']} is about "
                f"{p['nearest_town_km']:.0f} km.", space_after_pt=5, space_before_pt=3)
    normal_para(doc, "Road access to the project site is available via the existing state highway network. "
                "Internal access roads will be constructed during the project development phase to connect the "
                "main road to the various array locations. These roads will be designed to support the movement "
                "of heavy equipment during construction and for routine maintenance vehicles during operations.",
                space_after_pt=5, space_before_pt=3)
    normal_para(doc, f"The nearest grid substation is located approximately {p['tneb_distance_km']:.1f} km from "
                "the project site, which will serve as the point of grid injection. A 33 kV overhead/underground "
                "transmission line will be constructed to connect the project's HT switchyard to the substation "
                "for evacuation of generated power.", space_after_pt=6, space_before_pt=3)

    h2(doc, "Solar Resource Assessment")
    normal_para(doc, "A comprehensive solar resource assessment has been carried out for the proposed project site "
                "using satellite-derived irradiance data from internationally recognised databases including "
                "Meteonorm and Solargis. The assessment confirms that the site receives solar radiation adequate "
                "for commercially viable solar power generation. The detailed PVSyst simulation report is "
                "provided in Annexure – A at the end of this report.", space_after_pt=10, space_before_pt=3)
    page_break(doc)

    # ── SECTION 6 – PROJECT AT A GLANCE ──────────────────────────────────────
    h1(doc, "6. PROJECT AT A GLANCE")
    make_table(doc,
        ["Sl.", "Parameter", "Details"],
        [
            ["",  "1.0  GENERAL", ""],
            ["1.1", "Project Developer / Promoter", f"{p['company_name']} ({p['company_short']})"],
            ["1.2", "The Project", f"{f2(cap_ac)} MW AC, {f2(cap_dc)} MWp DC Ground Mounted Solar PV"],
            ["1.3", "Location of Plant", f"{p['location_village']} Village, {p['location_taluk']}, {p['location_district']}, {p['location_state']}"],
            ["1.4", "District", p["location_district"]],
            ["1.5", "State", p["location_state"]],
            ["1.6", "Land Required", f"Approx. {p['land_acres']:.0f} Acres (Private Land)"],
            ["1.7", "COD", f"{p['cod_month']} {p['cod_year']}"],
            ["",  "2.0  TECHNICAL", ""],
            ["2.1", "Installed Capacity (DC)", f"{f2(cap_dc)} MWp (Monocrystalline)"],
            ["2.2", "Installed Capacity (AC)", f"{f2(cap_ac)} MW AC"],
            ["2.3", "AC:DC Ratio", f"1 : {p['ac_dc_ratio']:.1f}"],
            ["2.4", "Technology", "Monocrystalline Silicon PV Modules with String/Central Inverters"],
            ["2.5", "Annual Generation (Year 1)", f"{f2(gen_yr1)} Lac Units ({gen_yr1/10:.1f} Million kWh)"],
            ["2.6", "Panel Degradation", f"{p['degradation_yr1_pct']:.1f}% in Year 1; {p['degradation_yr2_pct']:.1f}% per year from Year 2 onwards"],
            ["2.7", "Date of Commissioning (COD)", f"{p['cod_month']} {p['cod_year']}"],
            ["2.8", "PPA Term", f"{p['ppa_term']} Years"],
            ["2.9", "PPA Tariff", f"₹ {p['ppa_tariff']:.2f} per unit (Flat for {p['ppa_term']} years)"],
            ["",  "3.0  FINANCIAL", ""],
            ["3.1", "Total Project Cost", f"₹ {p['project_cost_cr']:.2f} Crores (₹ {f2(p['project_cost_lac'])} Lakhs)"],
            ["3.2", "Debt", f"₹ {p['debt_cr']:.2f} Crores (₹ {f2(p['debt_lac'])} Lakhs) – {p['debt_pct']:.0f}% of Project Cost"],
            ["3.3", "Equity", f"₹ {p['equity_cr']:.2f} Crores (₹ {f2(p['equity_lac'])} Lakhs) – {100-p['debt_pct']:.0f}% of Project Cost"],
            ["3.4", "Debt Interest Rate", f"{p['debt_interest_rate']:.2f}% per annum"],
            ["3.5", "Debt Repayment Tenor", f"{p['debt_tenor_yrs']} Years from COD"],
            ["3.6", "Project IRR (Post-Tax)", f"{proj_irr:.2f}%"],
            ["3.7", "Equity IRR (Post-Tax)", f"{eq_irr:.2f}%"],
            ["3.8", "Minimum DSCR", f"{min_dscr:.2f}"],
            ["3.9", "Average DSCR", f"{avg_dscr:.2f}"],
            ["",  "4.0  O&M COSTS", ""],
            ["4.1", "O&M – Year 1 (post COD)", "Free of Cost (included in EPC warranty)"],
            ["4.2", "O&M – From Year 2 onwards", f"₹ {p['om_rate_lac_per_mwac']:.1f} Lakhs per MWac = ₹ {f2(p['om_rate_lac_per_mwac']*cap_ac)} Lakhs p.a. (Base Year 2)"],
            ["4.3", "O&M Escalation", f"{p['om_escalation_pct']:.1f}% per annum"],
            ["4.4", "Insurance", f"{p['insurance_pct']:.2f}% of Project Cost = ₹ {f2(p['project_cost_lac']*p['insurance_pct']/100)} Lakhs (Year 1); {p['insurance_esc_pct']:.1f}% escalation p.a."],
        ],
        col_widths_cm=[1.2, 5.8, 8.6], center_cols=[0])
    page_break(doc)

    # ── SECTION 7 – DEMAND ANALYSIS ──────────────────────────────────────────
    h1(doc, "7. DEMAND ANALYSIS AND JUSTIFICATION OF THE PROJECT")
    h2(doc, "1.0 Introduction")
    normal_para(doc, f"Electricity is the most essential input for growth and development of any state. "
                f"{p['location_state']} is planning to grow rapidly in both the industrial and agricultural "
                "sectors and consequently the demand for power is on the rise. However, there exists a "
                "significant gap between conventional and non-conventional power generation capacity in the state.",
                space_after_pt=6, space_before_pt=3)
    normal_para(doc, f"The average electricity tariff for HT industrial consumers in {p['location_state']} "
                "currently ranges from ₹ 6.50 to ₹ 8.50 per unit (inclusive of fuel surcharge and other charges), "
                f"making solar PPAs at ₹ {p['ppa_tariff']:.2f} per unit economically compelling.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "2.0 Solar Power Potential in India")
    normal_para(doc, "India is endowed with rich solar energy resources. The average intensity of solar radiation "
                "received in India is 20 MW per km². Tamil Nadu receives solar energy with GHI values ranging "
                "from 5.0 to 6.0 kWh/m²/day, making it one of the most suitable states for large-scale solar "
                "deployment.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, "India achieved a cumulative solar capacity of over 70 GW as of March 2023. As part of the "
                "National Solar Mission, the Government of India aims to achieve 300 GW of solar capacity by 2030.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "3.0 World Energy Scenario")
    normal_para(doc, "The global transition toward renewable energy is accelerating. Solar energy is now the "
                "cheapest source of electricity in history. Over 29% of global electricity generation came from "
                "renewable sources in 2022. The global solar PV installed capacity crossed 1,000 GW (1 TW) in "
                "2022 – a milestone achieved faster than any other energy technology in history.",
                space_after_pt=6, space_before_pt=3)
    normal_para(doc, "Solar PV module prices have fallen by over 90% in the last decade, making solar the "
                "lowest-cost source of new electricity generation in most markets. In India, solar tariffs have "
                "fallen from over ₹ 7 per unit in 2014 to below ₹ 2.50 per unit in competitive auctions.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "4.0 National Action Plan on Climate Change (NAPCC)")
    normal_para(doc, "The National Action Plan on Climate Change (NAPCC) was launched by the Government of India "
                "in 2008 and identifies eight national missions for sustainable development. The National Solar "
                "Mission, launched in 2010, set a target of 100,000 MW (100 GW) of solar capacity. India is on "
                "track to achieve even more ambitious targets by 2030.",
                space_after_pt=6, space_before_pt=3)
    normal_para(doc, "The proposed project directly contributes to the objectives of the National Solar Mission "
                "and the NAPCC by adding clean solar generation capacity to the grid.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "5.0 Justification for the Project")
    normal_para(doc, "The proposed Solar Power Project is justified on the following grounds:",
                space_after_pt=4, space_before_pt=3)
    for item in [
        f"Economic Justification: Project IRR of {proj_irr:.2f}% (post-tax); flat tariff provides significant savings over escalating grid tariff expected to increase at 3–5% per year",
        f"Environmental Justification: ~{gen_yr1*0.82*100:,.0f} tonnes CO₂ avoided annually; over 25 years total CO₂ avoidance of approximately {gen_yr1*0.82*100*25/100000:.1f} lakh tonnes",
        f"Technical Justification: Excellent solar resource at {p['location_village']} – GHI 5.5–5.8 kWh/m²/day; flat shadow-free terrain ideal for ground mounting",
        f"Policy Justification: Aligns with {p['location_state']} Solar Energy Policy 2023, India's NDC commitments, and National Solar Mission objectives",
    ]:
        bullet(doc, item)
    page_break(doc)

    # ── SECTION 8 – BENEFITS ─────────────────────────────────────────────────
    h1(doc, "8. BENEFITS OF GRID-CONNECTED SOLAR PV POWER PLANT")
    h2(doc, "8.1 Economic Benefits")
    for item in [
        "Clean, renewable energy with zero fuel cost and zero emissions during operation",
        f"Long operational life of {p['ppa_term']}+ years with minimal maintenance",
        f"Predictable revenue stream through long-term PPA at fixed tariff of ₹ {p['ppa_tariff']:.2f} per unit",
        f"Supports India's renewable energy targets under National Solar Mission",
        "Reduces dependence on fossil fuel-based power and grid volatility",
        f"Project IRR {proj_irr:.2f}%; Equity IRR {eq_irr:.2f}% – strong investor returns",
        f"Significant savings for consumer vs HT grid tariff of ₹ 6.50–₹ 8.50 per unit",
    ]: bullet(doc, item)

    h2(doc, "8.2 Environmental Benefits")
    for item in [
        f"~{gen_yr1*0.82*100:,.0f} tonnes of CO₂ equivalent emissions avoided per year",
        f"Over 25 years, total CO₂ avoidance of approximately {gen_yr1*0.82*100*25/100000:.1f} lakh tonnes",
        "Zero water consumption for power generation",
        "No air, water, or noise pollution during operation – minimal environmental impact",
        "Supports India's Nationally Determined Contributions (NDCs) under the Paris Climate Agreement",
        f"Contributes to {p['location_state']}'s Renewable Purchase Obligation (RPO) compliance",
    ]: bullet(doc, item)

    h2(doc, "8.3 Technical Benefits")
    for item in [
        "Modular and scalable technology – easy to expand or modify",
        "High reliability with no moving parts in the PV modules – MTBF exceeds 20 years",
        "Grid stability support through modern inverters with reactive power compensation and LVRT",
        "SCADA-based monitoring enabling real-time performance tracking and remote diagnostics",
        "Short construction timeline of 4–6 months from financial close to commissioning",
        "Low water requirement – only panel washing",
    ]: bullet(doc, item)

    h2(doc, "8.4 Social and National Benefits")
    for item in [
        "Supports India's target of 500 GW of renewable energy by 2030",
        "Reduces India's dependence on fossil fuel imports, improving energy security",
        "~200 construction jobs; ~15 permanent O&M jobs",
        f"Local economic development in {p['location_district']} District through land lease and procurement",
        "Supports India's commitment under the Paris Agreement",
    ]: bullet(doc, item)
    page_break(doc)

    # ── SECTION 9 – SYSTEM DESCRIPTION ──────────────────────────────────────
    h1(doc, "9. BASIC SYSTEM DESCRIPTION")
    normal_para(doc, f"The proposed {f2(cap_dc)} MWp (DC) / {f2(cap_ac)} MW (AC) solar power plant will "
                "generate electricity from non-conventional solar sources using monocrystalline silicon "
                f"technology. The plant will use high-efficiency modules with an AC:DC ratio of "
                f"{p['ac_dc_ratio']:.1f}. The plant will be designed per IEC 61215, IEC 62109, and applicable "
                "CEA regulations.", space_after_pt=6, space_before_pt=3)

    h2(doc, "System Architecture")
    normal_para(doc, "The overall system follows a DC-optimised string configuration. Solar modules are arranged "
                "in strings, connected to string combiner boxes (SCB), which feed into string inverters. The AC "
                "output of multiple inverters is aggregated, stepped up by a transformer to 33 kV, and then "
                "evacuated to the grid substation. The plant is monitored through a central SCADA system.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "Main Components")
    for title, desc in [
        ("9.1 Solar PV Modules",
         f"High-efficiency monocrystalline solar PV modules (530–580 Wp per module) will be used. Total DC "
         f"capacity: {f2(cap_dc)} MWp. Total module count: approx. {cap_dc*1750:,.0f} modules. Modules will be "
         f"mounted on galvanised iron (GI) module mounting structures, oriented south at an optimal tilt angle "
         "of 11°–15°. Only Tier-1 rated modules with 25-year linear performance warranty will be used."),
        ("9.2 Inverters",
         f"String inverters or central inverters (ILR = {p['ac_dc_ratio']:.1f}) will be used. IEEE 1547/IEC 62109 "
         "certified with grid-support functions including reactive power control, frequency-watt control, and "
         "Low Voltage Ride-Through (LVRT). MPPT efficiency ≥99.5%; European efficiency ≥98.5%."),
        ("9.3 Module Mounting Structures",
         "Hot-dip galvanised GI structures designed for a minimum 25-year structural life. Designed to withstand "
         "wind speeds of up to 170 km/h and seismic loads as per IS 875 Part 3."),
        ("9.4 Transformers",
         f"Step-up transformers (0.69 kV / 33 kV) at each inverter block. Rated 5–6.3 MVA (ONAN type). "
         "Auxiliary transformers of 15 kVA capacity for station auxiliary supply."),
        ("9.5 High-Tension (HT) Switchyard",
         "33 kV HT switchyard with VCB panels, metering cubicles (0.2 class), and protection relays. "
         "DC battery backup for control and protection systems. 33 kV evacuation line to grid substation "
         f"(~{p['tneb_distance_km']:.1f} km)."),
        ("9.6 SCADA & Monitoring System",
         "Comprehensive SCADA-based plant monitoring. Real-time monitoring of generation, inverters, weather, "
         "and energy meters. Modbus TCP/IP and IEC 61850 communication. Daily/weekly/monthly performance reports."),
        ("9.7 Earthing & Lightning Protection",
         "GI flat earthing grid covering the entire plant. Lightning arrestors at module arrays, inverter "
         "stations, and HT switchyard."),
        ("9.8 Cables",
         "1500 VDC UV-resistant TÜV/UL certified solar DC cables. XLPE aluminium armoured 33 kV HT cables "
         "for grid connection."),
    ]:
        h3(doc, title)
        normal_para(doc, desc, space_after_pt=6, space_before_pt=3)
    page_break(doc)

    # ── SECTION 10 – BOQ ─────────────────────────────────────────────────────
    h1(doc, "10. BILL OF QUANTITY (BOQ)")
    make_table(doc,
        ["Sl.No.", "Description", "Unit", "Qty", "Remarks"],
        [
            ["",  "A. SOLAR PV MODULES", "", "", ""],
            ["1", f"Monocrystalline Solar PV Modules (530–580 Wp)", "Nos", f"~{cap_dc*1750:,.0f}", f"{f2(cap_dc)} MWp total"],
            ["",  "B. INVERTERS & POWER CONVERSION", "", "", ""],
            ["2", "String / Central Inverters (appropriate capacity)", "Nos", "LS", f"{f2(cap_ac)} MW AC"],
            ["3", "ACDB / DCDB Panels", "Set", "LS", ""],
            ["",  "C. MODULE MOUNTING STRUCTURES", "", "", ""],
            ["4", "GI Hot-Dip Galvanized Module Mounting Structures", "MT", "LS", "Fixed tilt"],
            ["5", "Foundation bolts, fasteners & accessories", "Lot", "1", ""],
            ["",  "D. HT INFRASTRUCTURE", "", "", ""],
            ["6", "33/0.69 kV Step-Up Transformer (solar duty)", "Nos", "4", "5–6.3 MVA each"],
            ["7", "33 kV HT Panel with 1600A ACB", "Nos", "1", ""],
            ["8", "VCB Panel (incoming + outgoing)", "Set", "1", "33 kV"],
            ["9", "33 kV Metering Cubicle (as per EB specs)", "Nos", "1", ""],
            ["10", "Station Battery & Float Cum Boost Charger", "Set", "1", ""],
            ["",  "E. CABLES", "", "", ""],
            ["11", "DC Cables from modules to combiner box & inverters", "Lot", "1", ""],
            ["12", "AC Cables from inverters to transformers", "Lot", "1", ""],
            ["13", f"33 kV XLPE Al Armoured Cable (evac. line ~{p['tneb_distance_km']:.1f} km)", "Lot", "1", ""],
            ["14", "Communication cables for SCADA", "Lot", "1", ""],
            ["",  "F. CIVIL WORKS", "", "", ""],
            ["15", f"Site preparation & levelling (~{p['land_acres']:.0f} acres)", "Lot", "1", ""],
            ["16", "Control & inverter room construction", "Nos", "1", ""],
            ["17", "Cable trenching", "Lot", "1", ""],
            ["18", "Internal roads, fencing & boundary wall", "Lot", "1", ""],
            ["19", "Water storage facility for panel cleaning", "Lot", "1", ""],
            ["",  "G. EARTHING & LIGHTNING PROTECTION", "", "", ""],
            ["20", "Earthing system (GI flat, earth pits, risers)", "Lot", "1", ""],
            ["21", "Lightning arrestors", "Lot", "1", ""],
            ["",  "H. SCADA, MONITORING & MISCELLANEOUS", "", "", ""],
            ["22", "SCADA & Plant Monitoring System", "Set", "1", ""],
            ["23", "Weather monitoring station", "Nos", "1", ""],
            ["24", "Firefighting system & safety equipment", "Lot", "1", ""],
        ],
        col_widths_cm=[1.0, 7.8, 1.3, 1.5, 3.0], center_cols=[0,2,3])
    normal_para(doc, "* This is a tentative BOM. The actual BOM will be finalised only after detailed engineering. "
                "Items and makes provided are indicative and subject to availability and pricing at the time of order.",
                italic=True, size_pt=9, color=RGBColor(0x66,0x66,0x66), space_after_pt=6)
    page_break(doc)

    # ── SECTION 11 – SCHEDULE ────────────────────────────────────────────────
    h1(doc, "11. PLANNED PROJECT SCHEDULE")
    make_table(doc,
        ["Project Activity", "Month 1", "Month 2", "Month 3", "Month 4"],
        [
            ["Detailed Survey of Land & Topography", "✓", "", "", ""],
            ["Finalization of Plant Structure & Design", "✓", "✓", "", ""],
            ["Preparation of Detailed Implementation Plan", "✓", "", "", ""],
            ["Firming up EPC Contract & Equipment Orders", "", "✓", "", ""],
            ["Land Preparation & Levelling", "", "✓", "✓", ""],
            ["Control Buildings & Cable Trenching", "", "", "✓", "✓"],
            ["Module Mounting Structures Installation", "", "", "✓", "✓"],
            ["Solar PV Module Installation", "", "", "✓", "✓"],
            ["Inverter & HT Electrical Equipment Installation", "", "", "", "✓"],
            ["SCADA & Monitoring System Installation", "", "", "", "✓"],
            ["Testing, Commissioning & COD", "", "", "", "✓"],
        ],
        col_widths_cm=[7.5, 2.0, 2.0, 2.0, 2.0])
    page_break(doc)

    # ── SECTION 11B – REGULATORY APPROVALS ──────────────────────────────────
    h1(doc, "11B. REGULATORY APPROVALS & CLEARANCES")
    normal_para(doc, f"The development of a {f2(cap_ac)} MW solar power project in {p['location_state']} requires "
                "regulatory approvals and clearances from central and state government authorities. "
                f"{p['company_short']} will obtain all necessary statutory approvals before commencement of construction.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "A. Central Government Approvals")
    for item in [
        "Registration with Ministry of New & Renewable Energy (MNRE) as an eligible solar power project",
        "Approval from Central Electricity Authority (CEA) for grid connectivity (if applicable)",
        "Environmental clearance – solar projects of this scale are generally exempted from EIA notification",
    ]: bullet(doc, item)

    h2(doc, f"B. State Government Approvals ({p['location_state']})")
    for item in [
        f"State Nodal Agency – Registration and consent for solar power project development",
        f"State DISCOM/Transco – Power Evacuation Agreement (PEA) and grid connection approval at 33 kV",
        f"State Electricity Regulatory Commission – Open Access / wheeling approval (if applicable)",
        "District Collector / Revenue – Land use conversion approval (if applicable)",
        "State Pollution Control Board – Consent to Establish (CTE) and Consent to Operate (CTO)",
        "PWD / Highways Department – Approval for approach road and utility crossings",
        "Revenue Department – Land registration, encumbrance certificate, and title documents",
    ]: bullet(doc, item)

    h2(doc, "C. Other Approvals")
    for item in [
        f"Chief Electrical Inspector to Government (CEIG), {p['location_state']} – HT electrical installation approval",
        "33 kV Transmission Line Permission from the State DISCOM",
        "Building permit for Control Room and Inverter Room from local body",
        "Fire NOC from Fire and Rescue Services Department",
        f"Power Purchase Agreement (PPA) execution with the identified off-taker for {p['ppa_term']} years at ₹ {p['ppa_tariff']:.2f}/unit",
    ]: bullet(doc, item)

    h2(doc, "D. Approval Timeline")
    make_table(doc,
        ["Approval / Clearance", "Expected Timeline", "Authority"],
        [
            ["State Nodal Agency Registration", "Month 1–2", "State Agency"],
            ["Land Registration & Title", "Month 1–3", "District Revenue Dept."],
            ["Power Evacuation Agreement", "Month 2–4", "State DISCOM"],
            ["CEIG Approval (HT Electrical)", "Month 3–5", f"CEIG, {p['location_state']}"],
            ["Environmental Consent", "Month 2–4", "State Pollution Control Board"],
            ["33 kV Transmission Line Permission", "Month 3–6", "State DISCOM"],
            ["Building Permit (Control Room)", "Month 2–4", "Local Body"],
            ["PPA Execution", "Month 1–3", f"Off-taker & {p['company_short']}"],
            ["Financial Close", "Month 4–7", f"Lenders & {p['company_short']}"],
            ["EPC Contract Award", "Month 6–8", p["company_short"]],
            ["COD (Commissioning)", f"{p['cod_month']} {p['cod_year']}", p["company_short"]],
        ],
        col_widths_cm=[7.0, 3.5, 5.0])
    page_break(doc)

    # ── SECTION 11A – O&M ────────────────────────────────────────────────────
    h1(doc, "11A. OPERATION & MAINTENANCE PLAN")
    normal_para(doc, f"The Operation and Maintenance (O&M) of the solar power plant is critical to ensuring "
                "optimal performance, long equipment life, and maximum energy generation over the "
                f"{p['ppa_term']}-year project life.", space_after_pt=6, space_before_pt=3)

    h2(doc, "O&M Cost Structure")
    for item in [
        "Year 1 (post-COD): O&M services are provided free of cost as part of the EPC contractor's warranty obligation.",
        f"Year 2 onwards: ₹ {p['om_rate_lac_per_mwac']:.1f} Lakhs/MWac = ₹ {f2(p['om_rate_lac_per_mwac']*cap_ac)} Lakhs/year (base Year 2); {p['om_escalation_pct']:.1f}% escalation p.a.",
        f"Insurance: {p['insurance_pct']:.2f}% of project cost = ₹ {f2(p['project_cost_lac']*p['insurance_pct']/100)} Lakhs (Year 1); {p['insurance_esc_pct']:.1f}% escalation p.a.",
    ]: bullet(doc, item)

    h2(doc, "O&M Scope of Services")
    for item in [
        "Preventive Maintenance: Scheduled inspection, cleaning, and servicing of all plant components",
        "Panel Cleaning: Regular module cleaning to maintain Performance Ratio – typically monthly or bi-monthly",
        "Corrective Maintenance: Target repair time <4 hours (critical equipment), <24 hours (others)",
        "Performance Monitoring: Continuous real-time SCADA monitoring; deviations >5% investigated immediately",
        "Vegetation Control: Regular clearance of weeds around structures to prevent shading and fire hazard",
        "Security: 24×7 security and surveillance including CCTV cameras, perimeter fencing, and security personnel",
        "Annual Maintenance Contracts (AMC) with OEMs for inverters, transformers, and SCADA systems",
    ]: bullet(doc, item)

    h2(doc, "Risk Analysis and Mitigation")
    make_table(doc,
        ["Risk Category", "Risk Description", "Mitigation Measure"],
        [
            ["Revenue Risk", "Solar resource variability", "P50 estimate; conservative degradation; 6-month DSRA maintained"],
            ["Off-take Risk", "PPA counterparty default", "Creditworthy off-taker; lender step-in rights"],
            ["Construction Risk", "Cost/schedule overrun", "Fixed-price EPC; Tier-1 contractor; contingency in budget"],
            ["Equipment Risk", "Module/inverter failure", "Tier-1 supplier warranties; on-site spare; equipment insurance"],
            ["Interest Rate Risk", "Rate increase", f"Fixed rate {p['debt_interest_rate']:.1f}% for full {p['debt_tenor_yrs']}-year tenor"],
            ["Regulatory Risk", "Policy/grid charge changes", "Long-term PPA provides price certainty"],
            ["Natural Disaster", "Cyclone/flood/earthquake", "170 km/h wind-rated structures; comprehensive insurance"],
            ["O&M Cost Risk", "Higher-than-expected costs", f"{p['om_escalation_pct']:.1f}% escalation assumed; fixed-price O&M contract"],
        ],
        col_widths_cm=[3.5, 5.0, 7.0])
    page_break(doc)

    # ── SECTION 12 – PROJECT COSTING ─────────────────────────────────────────
    h1(doc, "12. PROJECT COSTING")
    pp2 = doc.add_paragraph()
    pp2.paragraph_format.space_after = Pt(8)
    r = pp2.add_run(f"A. {f2(cap_dc)} MWp DC / {f2(cap_ac)} MW AC System – Ground Mounted Structure with "
                    "Monocrystalline Cells & String/Central Inverters")
    r.bold = True; r.italic = True; r.font.size = Pt(11); r.font.name = "Arial"
    normal_para(doc, "Figures in INR Lakhs", bold=True, size_pt=10,
                space_after_pt=6, space_before_pt=2, align=WD_ALIGN_PARAGRAPH.LEFT)

    pc = p["project_cost_lac"]
    make_table(doc,
        ["Sl.No.", "Description", "Unit", "Qty", "GST %", "Basic (₹ Lac)", "GST (₹ Lac)", "Amt with GST (₹ Lac)"],
        [
            ["1", f"Supply of Solar PV Modules – {f2(cap_dc)} MWp (Monocrystalline)", "MWp", f"{f2(cap_dc)}", "5%",   f"{pc*0.332:.2f}", f"{pc*0.332*0.05:.2f}", f"{pc*0.332*1.05:.2f}"],
            ["2", "Supply & Inst. of Inverters (String/Central)", "Nos", "LS", "18%",  f"{pc*0.095:.2f}", f"{pc*0.095*0.18:.2f}", f"{pc*0.095*1.18:.2f}"],
            ["3", "Module Mounting Structures (GI Hot-Dip Galvanized)", "MT", "LS", "18%",  f"{pc*0.090:.2f}", f"{pc*0.090*0.18:.2f}", f"{pc*0.090*1.18:.2f}"],
            ["4", "HT/LT & DC Cables", "Lot", "1", "18%",  f"{pc*0.038:.2f}", f"{pc*0.038*0.18:.2f}", f"{pc*0.038*1.18:.2f}"],
            ["5", "Transformers (33/0.69 kV)", "Nos", "4", "18%",  f"{pc*0.038:.2f}", f"{pc*0.038*0.18:.2f}", f"{pc*0.038*1.18:.2f}"],
            ["6", "HT Switchyard, VCB & Metering", "Set", "1", "18%",  f"{pc*0.029:.2f}", f"{pc*0.029*0.18:.2f}", f"{pc*0.029*1.18:.2f}"],
            ["7", "Civil Works (Foundation, Roads, Fencing)", "Lot", "1", "18%",  f"{pc*0.063:.2f}", f"{pc*0.063*0.18:.2f}", f"{pc*0.063*1.18:.2f}"],
            ["8", "SCADA & Monitoring", "Set", "1", "18%",  f"{pc*0.014:.2f}", f"{pc*0.014*0.18:.2f}", f"{pc*0.014*1.18:.2f}"],
            ["9", "Earthing & Lightning Protection", "Lot", "1", "18%",  f"{pc*0.010:.2f}", f"{pc*0.010*0.18:.2f}", f"{pc*0.010*1.18:.2f}"],
            ["10", f"Evacuation Line ({p['tneb_distance_km']:.1f} km, 33 kV)", "Km", f"{p['tneb_distance_km']:.1f}", "18%",  f"{pc*0.024:.2f}", f"{pc*0.024*0.18:.2f}", f"{pc*0.024*1.18:.2f}"],
            ["11", "Land Development & Site Fencing", "Lot", "1", "18%",  f"{pc*0.016:.2f}", f"{pc*0.016*0.18:.2f}", f"{pc*0.016*1.18:.2f}"],
            ["12", "Transportation, Erection & Commissioning", "Lot", "1", "18%",  f"{pc*0.040:.2f}", f"{pc*0.040*0.18:.2f}", f"{pc*0.040*1.18:.2f}"],
            ["13", "Pre-operative Expenses, IDC & Margin Money", "Lot", "1", "N/A", f"{pc*0.051:.2f}", "–", f"{pc*0.051:.2f}"],
            ["", "TOTAL PROJECT COST", "", "", "", "–", "–", f"{pc:.2f}"],
        ],
        col_widths_cm=[0.7, 5.5, 1.0, 0.9, 0.9, 2.2, 2.0, 2.4], header_size=8, body_size=8)

    normal_para(doc, f"* The total project cost is rounded to INR {pc:.0f} Lakhs (₹ {p['project_cost_cr']:.2f} Crores) "
                "inclusive of all capital expenditure, pre-operative expenses, and interest during construction.",
                italic=True, size_pt=9, color=RGBColor(0x55,0x55,0x55), space_after_pt=8)

    h2(doc, "Financing Structure")
    make_table(doc,
        ["Component", "Amount (₹ Lac)", "Percentage"],
        [
            ["Total Project Cost",            f"{f2(p['project_cost_lac'])}", "100%"],
            [f"Debt (Term Loan @ {p['debt_interest_rate']:.1f}%)", f"{f2(p['debt_lac'])}", f"{p['debt_pct']:.0f}%"],
            ["Equity (by " + p['company_short'] + ")", f"{f2(p['equity_lac'])}", f"{100-p['debt_pct']:.0f}%"],
        ],
        col_widths_cm=[6.0, 3.5, 2.0])
    page_break(doc)

    # ── SECTION 13 – FINANCIALS + 25-YR CASH FLOW ────────────────────────────
    h1(doc, "13. PROJECT FINANCIALS – 25 YEAR CASH FLOW (₹ in Lakhs)")
    h2(doc, "Key Assumptions")
    make_table(doc,
        ["Parameter", "Assumption"],
        [
            ["Installed Capacity (AC / DC)", f"{f2(cap_ac)} MW AC / {f2(cap_dc)} MWp DC"],
            ["Year 1 Generation", f"{f2(gen_yr1)} Lac Units"],
            ["Degradation – Year 1 / Year 2+", f"{p['degradation_yr1_pct']:.1f}% / {p['degradation_yr2_pct']:.1f}% p.a."],
            ["PPA Tariff", f"₹ {p['ppa_tariff']:.2f} per kWh (Flat, {p['ppa_term']} years)"],
            ["O&M – Year 1", "Nil (Free – EPC warranty)"],
            ["O&M – Year 2 base", f"₹ {f2(p['om_rate_lac_per_mwac']*cap_ac)} Lakhs; {p['om_escalation_pct']:.1f}% escalation p.a."],
            ["Insurance", f"{p['insurance_pct']:.2f}% of cost p.a.; {p['insurance_esc_pct']:.1f}% escalation p.a."],
            ["Book Depreciation (SLM)", f"{p['ppa_term']} years = ₹ {f2(dep_slm)} Lakhs/year"],
            ["Tax Depreciation", "WDV @ 40% p.a."],
            ["Tax Regime", "Old Regime with MAT (effective MAT ~16.69%; normal ~33.38%)"],
            ["Debt", f"₹ {f2(p['debt_lac'])} Lakhs @ {p['debt_interest_rate']:.1f}% p.a."],
            ["Debt Tenor / Annual Principal", f"{p['debt_tenor_yrs']} Years / ₹ {f2(ann_prin)} Lakhs"],
            ["Equity", f"₹ {f2(p['equity_lac'])} Lakhs (contributed by {p['company_short']})"],
        ],
        col_widths_cm=[7.0, 8.5])

    # 25-year cash flow table
    normal_para(doc, "25-Year Projected Cash Flow Statement (₹ in Lakhs)",
                bold=True, size_pt=11, color=C_NAVY,
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=5, space_before_pt=8)

    cf_hdrs = ["Yr", "Gen\n(Lac U)", "Revenue", "O&M", "Insur.", "EBITDA",
               "Depn.", "EBIT", "Interest", "PBT", "Tax", "PAT", "Principal", "DSCR"]
    cf_wids = [0.6, 1.5, 1.7, 1.5, 1.4, 1.7, 1.3, 1.6, 1.5, 1.6, 1.4, 1.5, 1.5, 1.2]

    table = doc.add_table(rows=1+len(cfs), cols=len(cf_hdrs))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(cf_hdrs):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "1F3864"); set_cell_border(cell, "888888")
        set_col_width(cell, cf_wids[i]); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = cell.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(h); r2.bold = True
        r2.font.size = Pt(7.5); r2.font.name = "Arial"; r2.font.color.rgb = C_WHITE

    for ri, cf in enumerate(cfs):
        bg = "EBF3FB" if ri % 2 == 1 else "FFFFFF"
        dscr_str = f2(cf["dscr"]) if cf["year"] <= p["debt_tenor_yrs"] else "N/A"
        vals = [str(cf["year"]), f2(cf["gen"]), f2(cf["revenue"]),
                f2(cf["om"]), f2(cf["insurance"]), f2(cf["ebitda"]),
                f2(cf["depreciation"]), f2(cf["ebit"] if "ebit" in cf else cf["ebitda"]-cf["depreciation"]),
                f2(cf["interest"]), f2(cf["pbt"]), f2(cf["tax"]),
                f2(cf["pat"]), f2(cf["principal"]), dscr_str]
        for ci, val in enumerate(vals):
            cell = table.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg); set_cell_border(cell, "CCCCCC")
            set_col_width(cell, cf_wids[ci]); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp2 = cell.paragraphs[0]; pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp2.paragraph_format.space_before = Pt(1); pp2.paragraph_format.space_after = Pt(1)
            run = pp2.add_run(val)
            run.font.size = Pt(7.5); run.font.name = "Arial"
            if ci == 13 and cf["year"] <= p["debt_tenor_yrs"]:
                try:
                    run.font.color.rgb = C_GREEN if cf["dscr"] >= 1.25 else C_RED
                    run.bold = True
                except: pass

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    normal_para(doc, "Note: DSCR = (EBITDA – Tax) / (Principal + Interest). "
                "Figures in ₹ Lakhs. Generation in Lac Units.",
                italic=True, size_pt=8, color=RGBColor(0x55,0x55,0x55), space_after_pt=6)
    page_break(doc)

    # ── SECTION 14 – DEBT & DSCR ─────────────────────────────────────────────
    h1(doc, "14. DEBT REPAYMENT SCHEDULE & DSCR ANALYSIS")
    normal_para(doc, f"The project is financed with term debt of ₹ {f2(p['debt_lac'])} Lakhs at "
                f"{p['debt_interest_rate']:.1f}% per annum, repayable over {p['debt_tenor_yrs']} years from COD "
                f"on a straight-line basis. Annual principal repayment: ₹ {f2(ann_prin)} Lakhs.",
                space_after_pt=8, space_before_pt=3)

    debt_rows = []
    for cf in cfs:
        if cf["year"] > p["debt_tenor_yrs"]: break
        debt_rows.append([str(cf["year"]), f2(cf["op_debt"]), f2(cf["principal"]),
                          f2(cf["interest"]), f2(cf["debt_service"]), f2(cf["cl_debt"]),
                          f2(cf["dscr"])])
    make_table(doc,
        ["Year", "Opening Debt\n(₹ Lac)", "Principal\n(₹ Lac)", "Interest\n(₹ Lac)",
         "Total Debt Svc\n(₹ Lac)", "Closing Debt\n(₹ Lac)", "DSCR"],
        debt_rows, col_widths_cm=[1.2, 2.8, 2.4, 2.4, 2.8, 2.8, 1.6])

    make_table(doc,
        ["DSCR Metric", "Value"],
        [
            ["Minimum DSCR (over loan tenor)", f"{min_dscr:.2f}x"],
            ["Average DSCR (over loan tenor)", f"{avg_dscr:.2f}x"],
        ],
        col_widths_cm=[9.0, 4.0], header_bg="2E74B5")
    normal_para(doc, f"DSCR Analysis: The project maintains a healthy DSCR throughout the loan tenor. "
                f"The minimum DSCR of {min_dscr:.2f}x is well above the threshold of 1.10x typically required "
                f"by lenders. The average DSCR of {avg_dscr:.2f}x demonstrates comfortable debt serviceability.",
                space_after_pt=10, space_before_pt=3)
    page_break(doc)

    # ── SECTION 15 – SUMMARY ─────────────────────────────────────────────────
    h1(doc, "15. SUMMARY OF RESULTS")
    # Big summary table matching template T13
    normal_para(doc, f"FINANCIAL SUMMARY – {f2(cap_ac)} MW SOLAR PROJECT – {p['company_short']}",
                bold=True, size_pt=11, color=C_NAVY,
                align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=5)
    make_table(doc,
        ["Parameter", "Value"],
        [
            ["Project Capacity (AC / DC)", f"{f2(cap_ac)} MW AC / {f2(cap_dc)} MWp DC"],
            ["Location", loc_str],
            ["COD", f"{p['cod_month']} {p['cod_year']}"],
            ["Total Project Cost", f"₹ {p['project_cost_cr']:.2f} Crores (₹ {f2(p['project_cost_lac'])} Lakhs)"],
            ["Debt", f"₹ {p['debt_cr']:.2f} Crores @ {p['debt_interest_rate']:.1f}% p.a. ({p['debt_tenor_yrs']} yrs)"],
            ["Equity", f"₹ {p['equity_cr']:.2f} Crores"],
            ["Year 1 Generation", f"{f2(gen_yr1)} Lac Units"],
            ["Year 1 Revenue", f"₹ {f2(cfs[0]['revenue'])} Lakhs"],
            ["Year 1 EBITDA", f"₹ {f2(cfs[0]['ebitda'])} Lakhs"],
            ["PPA Tariff", f"₹ {p['ppa_tariff']:.2f}/unit (Flat, {p['ppa_term']} years)"],
            ["Project IRR (Post-Tax)", f"{proj_irr:.2f}%"],
            ["Equity IRR (Post-Tax)", f"{eq_irr:.2f}%"],
            ["Minimum DSCR", f"{min_dscr:.2f}x"],
            ["Average DSCR", f"{avg_dscr:.2f}x"],
            ["Specific Cost", f"₹ {p['project_cost_cr']/cap_ac:.2f} Crores/MWac"],
        ],
        col_widths_cm=[7.5, 8.1])
    page_break(doc)

    # ── SECTION 16 – CONCLUSION ──────────────────────────────────────────────
    h1(doc, "16. CONCLUSION")
    normal_para(doc, f"The detailed financial analysis for the {f2(cap_ac)} MW AC ({f2(cap_dc)} MWp DC) "
                f"ground-mounted solar power project by M/s. {p['company_name']} ({p['company_short']}) at "
                f"{p['location_village']}, {p['location_district']} District, {p['location_state']}, "
                "demonstrates that the project is both technically feasible and financially attractive.",
                space_after_pt=6, space_before_pt=3)

    h2(doc, "Key Findings:")
    for item in [
        f"The project achieves a healthy Project IRR of {proj_irr:.2f}% (post-tax) over a {p['ppa_term']}-year PPA term, well above the weighted average cost of capital.",
        f"Equity IRR of {eq_irr:.2f}% demonstrates strong return on equity investment for {p['company_short']} shareholders.",
        f"The minimum DSCR of {min_dscr:.2f}x (above the standard lender threshold of 1.10x) confirms comfortable debt serviceability throughout the {p['debt_tenor_yrs']}-year loan tenor.",
        f"The average DSCR of {avg_dscr:.2f}x over the debt repayment period provides significant comfort to lenders.",
        f"The flat PPA tariff of ₹ {p['ppa_tariff']:.2f} per unit for {p['ppa_term']} years ensures revenue certainty and protects against market price volatility.",
        f"The AC:DC ratio of {p['ac_dc_ratio']:.1f} optimises generation and reduces per-unit cost of electricity.",
        f"O&M costs are managed efficiently with Year 1 free of charge and {p['om_escalation_pct']:.1f}% annual escalation from Year 2.",
        f"The project will generate approximately {f2(gen_yr1)} Lac Units in Year 1, contributing to {p['location_state']}'s renewable energy targets.",
    ]: bullet(doc, item)

    h2(doc, "Overall Assessment")
    normal_para(doc, "The project is technically sound, with established monocrystalline PV technology, proven "
                "equipment suppliers, and straightforward grid integration at 33 kV. The site characteristics "
                "including flat shadow-free terrain, good solar resource, proximity to grid infrastructure, and "
                "availability of private land are all highly favourable.", space_after_pt=6, space_before_pt=3)
    normal_para(doc, "As seen from the above financial summary, the project has a healthy DSCR for the entire "
                "duration of the debt. Hence, debt servicing is feasible. Apart from this, the project also "
                "provides a steady and attractive return on investment for the equity employed.",
                space_after_pt=6, space_before_pt=3)
    normal_para(doc, "With the above assumptions and financial results, it is clear that the project is "
                "Technically and Financially viable and is worth investing. The project is recommended for "
                "approval and funding.", bold=True, space_after_pt=16, space_before_pt=6)

    for line, sz, bold, col in [
        ("***************************", 12, False, C_NAVY),
        (p["company_name"].upper(), 13, True, C_NAVY),
        (p["company_address"], 10, False, RGBColor(0x44,0x44,0x44)),
        (f"A Subsidiary of {ownership_str(p)}", 10, False, RGBColor(0x55,0x55,0x55)),
    ]:
        pp3 = doc.add_paragraph()
        pp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp3.paragraph_format.space_after = Pt(4)
        r3 = pp3.add_run(line); r3.bold = bold
        r3.font.size = Pt(sz); r3.font.name = "Arial"; r3.font.color.rgb = col
    page_break(doc)

    # ── ANNEXURE A – PVSYST ──────────────────────────────────────────────────
    h1(doc, "ANNEXURE – A: SOLAR RESOURCE ASSESSMENT & PVSYST SIMULATION")
    normal_para(doc, f"This Annexure contains the solar resource assessment data and PVSyst energy yield "
                f"simulation results for the {f2(cap_ac)} MW AC project at {p['location_village']}, "
                f"{p['location_district']} District, {p['location_state']}.",
                space_after_pt=8, space_before_pt=3)

    h2(doc, "A.1 Solar Resource Data Summary")
    normal_para(doc, "Monthly Average Global Horizontal Irradiance (GHI) at project site (kWh/m²/day):",
                space_after_pt=4, space_before_pt=3)
    make_table(doc,
        ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"],
        [["5.80","6.10","6.40","6.20","5.90","4.80","4.50","4.70","5.20","5.60","5.40","5.50"]],
        col_widths_cm=[1.3]*12, body_size=9)

    h2(doc, "A.2 PVSyst Simulation Parameters")
    make_table(doc,
        ["PVSyst Parameter", "Value Used"],
        [
            ["Simulation Software",         "PVSyst Version 7.x"],
            ["Solar Resource Database",     "Meteonorm / Solargis"],
            ["Module Type",                 "Monocrystalline Si – 550–580 Wp"],
            ["Module Efficiency",           "21.5% – 22.5%"],
            ["Tilt Angle",                  f"12°–15° (Fixed South-facing)"],
            ["DC/AC Ratio (ILR)",           f"{p['ac_dc_ratio']:.1f}"],
            ["Soiling Loss",                "2.0% per annum"],
            ["Mismatch Loss",               "2.0%"],
            ["DC Wiring Loss",              "1.5%"],
            ["Transformer Loss",            "1.0%"],
            ["System Availability",         "98.5%"],
            ["Performance Ratio (PR)",      "~78% – 80%"],
            ["Specific Yield",              f"~{p['gen_yr1_lac_per_mwac']*100/cap_dc:.0f} kWh/kWp/year"],
            ["P50 Annual Generation (Yr 1)", f"~{f2(gen_yr1)} Lac Units"],
            ["P90 Annual Generation (Yr 1)", f"~{gen_yr1*0.953:.0f} Lac Units"],
        ],
        col_widths_cm=[8.0, 7.5])

    h2(doc, "A.3 Month-wise Generation Estimate")
    monthly_pct = [8.95,8.09,9.05,8.72,8.65,6.77,6.70,7.02,7.53,8.47,7.86,8.47]
    mgen_rows = [[m, f"{gen_yr1*pct/100:.2f}", f"{pct:.1f}%"]
                 for m, pct in zip(["January","February","March","April","May","June",
                                    "July","August","September","October","November","December"],
                                   monthly_pct)]
    mgen_rows.append(["Annual Total", f"{gen_yr1:.2f}", "100%"])
    make_table(doc, ["Month", "Generation (Lac Units)", "% of Annual"],
               mgen_rows, col_widths_cm=[3.5, 4.5, 3.0])
    normal_para(doc, "Note: Monthly estimates are illustrative based on typical solar profile. "
                "Replace with actual PVSyst simulation report upon completion of site-specific assessment.",
                italic=True, size_pt=9, color=RGBColor(0x55,0x55,0x55), space_after_pt=8)

    h2(doc, "A.4 Energy Generation Summary – 25 Years")
    gen25_rows = [[str(cf["year"]), f"{cf['gen']:.2f}", f"{cf['revenue']:.2f}"] for cf in cfs]
    make_table(doc, ["Year", "Generation (Lac Units)", "Revenue (₹ Lakhs)"],
               gen25_rows, col_widths_cm=[2.0, 4.5, 4.5])

    pp_end = doc.add_paragraph()
    pp_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp_end.paragraph_format.space_before = Pt(20)
    r_end = pp_end.add_run("– End of Detailed Project Report –")
    r_end.bold = True; r_end.italic = True; r_end.font.size = Pt(13)
    r_end.font.color.rgb = C_NAVY; r_end.font.name = "Arial"

    # ── Save ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
